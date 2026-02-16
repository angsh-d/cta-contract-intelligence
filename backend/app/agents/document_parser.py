"""Tier 1: DocumentParserAgent — PDF/DOCX text/table extraction and LLM structuring."""

import asyncio
import logging
import os
import time
from typing import Any, Optional
from app.agents.base import BaseAgent
from app.agents.config import AgentConfig
from app.exceptions import LLMResponseError
from app.models.agent_schemas import (
    DocumentMetadata, DocumentParseInput, DocumentParseOutput,
    ParsedSection, ParsedTable,
)

logger = logging.getLogger(__name__)


class DocumentParserAgent(BaseAgent):
    """Extract text, structure, tables, and metadata from PDFs and DOCX files."""

    def __init__(self, config, llm_provider, prompt_loader, vector_store,
                 progress_callback=None, fallback_provider=None,
                 trace_context=None, llm_semaphore=None):
        super().__init__(config, llm_provider, prompt_loader, progress_callback,
                         fallback_provider, trace_context, llm_semaphore)
        self.vector_store = vector_store

    async def process(self, input_data: DocumentParseInput) -> DocumentParseOutput:
        start = time.monotonic()

        # Step 1: Extract raw text (blocking I/O — run in thread)
        raw_text, page_count = await asyncio.to_thread(
            self._extract_text, input_data.file_path
        )
        await self._report_progress("text_extraction", 20, f"Extracted {len(raw_text)} chars from {page_count} pages")

        # Step 2: Extract tables (blocking I/O — run in thread)
        tables = await asyncio.to_thread(
            self._extract_tables, input_data.file_path
        )
        await self._report_progress("table_extraction", 35, f"Found {len(tables)} tables")

        # Step 3: Chunk if needed (50K threshold to avoid LLM output truncation)
        chunks = self._chunk_text(raw_text) if len(raw_text) > 50_000 else [raw_text]

        # Step 4: LLM structuring (per chunk)
        all_sections: list[ParsedSection] = []
        metadata: Optional[DocumentMetadata] = None
        result = {}

        for i, chunk in enumerate(chunks):
            system_prompt = self.prompts.get("document_parser_system")
            user_prompt = self.prompts.get(
                "document_parser_extraction",
                document_type=input_data.document_type.value,
                raw_text=chunk,
            )
            result = await self.call_llm(system_prompt, user_prompt)
            sections_raw = result.get("sections")
            if not sections_raw:
                raise LLMResponseError(f"DocumentParserAgent: LLM response missing 'sections' for chunk {i+1}")
            chunk_sections = []
            for s in sections_raw:
                # Coerce null strings to "" for required fields, strip nulls for optional/defaulted fields
                cleaned = {}
                for k, v in s.items():
                    if v is None and k in ("section_number", "section_title", "text"):
                        cleaned[k] = ""
                    elif v is not None:
                        cleaned[k] = v
                chunk_sections.append(ParsedSection(**cleaned))
            all_sections.extend(chunk_sections)

            if metadata is None and result.get("metadata"):
                meta_raw = result["metadata"]
                # Coerce amendment_number to int if LLM returned a word (e.g. "Second")
                if "amendment_number" in meta_raw and meta_raw["amendment_number"] is not None:
                    try:
                        meta_raw["amendment_number"] = int(meta_raw["amendment_number"])
                    except (ValueError, TypeError):
                        meta_raw["amendment_number"] = None
                metadata = DocumentMetadata(**{k: v for k, v in meta_raw.items() if v is not None})

            await self._report_progress(
                "llm_structuring",
                40 + int(40 * (i + 1) / len(chunks)),
                f"Structured chunk {i+1}/{len(chunks)}",
            )

        # Step 5: Deduplicate sections from overlapping chunks
        sections = self._deduplicate_sections(all_sections)

        # Step 6: Embed and upsert to pgvector
        await self._upsert_embeddings(input_data.contract_stack_id, sections, metadata, input_data.document_id)
        await self._report_progress("embedding", 95, "Embeddings upserted to pgvector")

        return DocumentParseOutput(
            document_id=input_data.document_id,
            metadata=metadata,
            sections=sections,
            tables=tables,
            raw_text=raw_text,
            char_count=len(raw_text),
            page_count=page_count,
            extraction_model=self.config.model_override or "",
            extraction_latency_ms=int((time.monotonic() - start) * 1000),
            llm_reasoning=result.get("reasoning", ""),
            extraction_confidence=result.get("extraction_confidence", 0.9),
        )

    def _extract_text(self, file_path: str) -> tuple[str, int]:
        """Dispatch text extraction by file extension."""
        ext = os.path.splitext(file_path)[1].lower()
        if ext == ".docx":
            return self._extract_text_docx(file_path)
        return self._extract_text_pymupdf(file_path)

    def _extract_tables(self, file_path: str) -> list[ParsedTable]:
        """Dispatch table extraction by file extension."""
        ext = os.path.splitext(file_path)[1].lower()
        if ext == ".docx":
            return self._extract_tables_docx(file_path)
        return self._extract_tables_pdfplumber(file_path)

    def _extract_text_pymupdf(self, file_path: str) -> tuple[str, int]:
        import fitz
        doc = fitz.open(file_path)
        try:
            pages_text = []
            for page in doc:
                pages_text.append(page.get_text())
            return "\n".join(pages_text), len(doc)
        finally:
            doc.close()

    def _extract_tables_pdfplumber(self, file_path: str) -> list[ParsedTable]:
        import pdfplumber
        tables = []
        with pdfplumber.open(file_path) as pdf:
            for page_num, page in enumerate(pdf.pages, 1):
                for table_data in (page.extract_tables() or []):
                    if table_data and len(table_data) > 1:
                        headers = [str(h or "") for h in table_data[0]]
                        rows = [[str(c or "") for c in row] for row in table_data[1:]]
                        tables.append(ParsedTable(
                            table_id=f"table_{page_num}_{len(tables)}",
                            headers=headers,
                            rows=rows,
                            page_number=page_num,
                        ))
        return tables

    def _extract_text_docx(self, file_path: str) -> tuple[str, int]:
        """Extract text from a DOCX file using python-docx."""
        from docx import Document as DocxDocument
        doc = DocxDocument(file_path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        full_text = "\n".join(paragraphs)
        # Estimate page count (~3000 chars per page)
        page_count = max(1, len(full_text) // 3000)
        return full_text, page_count

    def _extract_tables_docx(self, file_path: str) -> list[ParsedTable]:
        """Extract tables from a DOCX file using python-docx."""
        from docx import Document as DocxDocument
        doc = DocxDocument(file_path)
        tables: list[ParsedTable] = []
        for idx, table in enumerate(doc.tables):
            rows_data = []
            for row in table.rows:
                rows_data.append([cell.text.strip() for cell in row.cells])
            if len(rows_data) > 1:
                headers = rows_data[0]
                rows = rows_data[1:]
                tables.append(ParsedTable(
                    table_id=f"table_docx_{idx}",
                    headers=headers,
                    rows=rows,
                    page_number=1,  # DOCX doesn't have page numbers
                ))
        return tables

    def _chunk_text(self, text: str, chunk_size: int = 50_000, overlap: int = 2_000) -> list[str]:
        chunks = []
        start = 0
        while start < len(text):
            end = start + chunk_size
            chunks.append(text[start:end])
            start = end - overlap
        return chunks

    def _deduplicate_sections(self, sections: list[ParsedSection]) -> list[ParsedSection]:
        seen: dict[str, ParsedSection] = {}
        for section in sections:
            if section.section_number not in seen:
                seen[section.section_number] = section
            else:
                if len(section.text) > len(seen[section.section_number].text):
                    seen[section.section_number] = section
        return list(seen.values())

    async def _upsert_embeddings(self, contract_stack_id, sections, metadata, document_id=None):
        """Upsert section embeddings to pgvector for semantic search."""
        if not sections:
            return
        section_dicts = [
            {
                "section_number": s.section_number,
                "section_title": s.section_title,
                "text": s.text,
            }
            for s in sections
        ]
        effective_date = str(metadata.effective_date) if metadata and metadata.effective_date else None
        await self.vector_store.upsert_embeddings(
            contract_stack_id=contract_stack_id,
            sections=section_dicts,
            document_id=document_id,
            effective_date=effective_date,
        )
